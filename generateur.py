import docx, random, copy

nom_du_fichier_doc = "test.docx"
nombre_de_copies = 3
verbes = [
    {
        "infinif": "cost",
        "preterit": "cost",
        "p.passe": "cost",
        "francais": "coûter"
    },
    {
        "infinif": "cut",
        "preterit": "cut",
        "p.passe": "cut",
        "francais": "couper"
    },
    {
        "infinif": "hit",
        "preterit": "hit",
        "p.passe": "hit",
        "francais": "frapper"
    },
    {
        "infinif": "hurt",
        "preterit": "hurt",
        "p.passe": "hurt",
        "francais": "blesser"
    },
    {
        "infinif": "let",
        "preterit": "let",
        "p.passe": "let",
        "francais": "permettre/louer"
    },
    {
        "infinif": "put",
        "preterit": "put",
        "p.passe": "put",
        "francais": "mettre"
    },
    {
        "infinif": "shut",
        "preterit": "shut",
        "p.passe": "shut",
        "francais": "fermer"
    },
    {
        "infinif": "begin",
        "preterit": "began",
        "p.passe": "begun",
        "francais": "commencer"
    },
    {
        "infinif": "drink",
        "preterit": "drank",
        "p.passe": "drunk",
        "francais": "boire"
    },
    {
        "infinif": "ring",
        "preterit": "rang",
        "p.passe": "rung",
        "francais": "sonner/téléphoner"
    },
    {
        "infinif": "sing",
        "preterit": "sang",
        "p.passe": "sung",
        "francais": "chanter"
    },
    {
        "infinif": "swim",
        "preterit": "swam",
        "p.passe": "swum",
        "francais": "nager"
    },
    {
        "infinif": "bleed",
        "preterit": "bled",
        "p.passe": "bled",
        "francais": "saigner"
    },
    {
        "infinif": "bring",
        "preterit": "brought",
        "p.passe": "brought",
        "francais": "apporter"
    },
    {
        "infinif": "build",
        "preterit": "built",
        "p.passe": "built",
        "francais": "nager"
    },
    {
        "infinif": "burn",
        "preterit": "burnt",
        "p.passe": "burnt",
        "francais": "brûler"
    },
    {
        "infinif": "buy",
        "preterit": "bought",
        "p.passe": "bought",
        "francais": "acheter"
    },
    {
        "infinif": "catch",
        "preterit": "caught",
        "p.passe": "caught",
        "francais": "attraper"
    },
    {
        "infinif": "dream",
        "preterit": "dreamt",
        "p.passe": "dreamt",
        "francais": "rêver"
    },
    {
        "infinif": "feel",
        "preterit": "felt",
        "p.passe": "felt",
        "francais": "sentir/ressentir"
    }, 
]

######################################
################ Code ################

doc = docx.Document()


for page in range(nombre_de_copies):
    # Titre principal
    doc.add_heading('Irregular verbs exam', 0)

    # Titre secondaire, premier exo
    doc.add_heading('Complete the array', level=1)
    doc.add_paragraph(
        'Complete this array with corrects irregular verbs or by translating them.'
    )

    table = doc.add_table(rows=1, cols=4)
    top_cells = table.rows[0].cells
    top_cells[0].text = "Verbe"
    top_cells[1].text = "Prétérit"
    top_cells[2].text = "P. passé"
    top_cells[3].text = "Traduction"

    copy_of_verbes = copy.deepcopy(verbes)
    random.shuffle(copy_of_verbes)
    for verbe in copy_of_verbes:
        row_cells = table.add_row().cells
        keys = list(verbe)
        rdm_index = random.choice(keys)
        for k in keys:
            if k != rdm_index:
                verbe[k] = ""
        row_cells[0].text = verbe["infinif"]
        row_cells[1].text = verbe["preterit"]
        row_cells[2].text = verbe["p.passe"]
        row_cells[3].text = verbe["francais"]

    doc.add_page_break()

doc.save(nom_du_fichier_doc)
